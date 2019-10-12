import csv
import random
from string import Template
from docx import Document
from docx.shared import Pt
from datetime import date


def get_file_name():
    while True:
        file_name = input("請輸入欲讀取csv檔名稱: ")
        if file_name.split(".")[-1] != "csv":
            print("請輸入.csv檔 (範例: sample.csv)")
            return True
        return file_name

def get_user_input():
    path_prefix = "./data/"
    file_list = []
    file_name = get_file_name()
    file_list.append(path_prefix + file_name)
    check_more_input = True
    while check_more_input:
        status = input("請問是否還要加入其他檔案? (y/N): ")
        print(status)
        if status == "y" or status == "Y" or status == "yes":
            file_name = get_file_name()
            file_list.append(path_prefix + file_name)
            print(file_list)
        elif status == "N" or status == "n" or status == "no":
            check_more_input = False
        else:
            print("請輸入 y 或 N ")
            check_more_input = True
            
    question_amount = input("請輸入題數: ")
    return file_list, question_amount

def load_data(path_list):
    english_list = []
    mandarin_list = []
    for path in path_list:
        with open(path, "r", encoding="UTF-8") as f:
            for line in f:
                try:
                    data = line.rstrip("\n").split(",")
                    if len(data) > 3:
                        english_list.append(data[1])
                        english_list.append(data[4])
                        mandarin_list.append(data[2])
                        mandarin_list.append(data[5])
                    else:
                        english_list.append(data[1])
                        mandarin_list.append(data[2])
                except:
                    print(line)
                    print("Parsing error")
                    pass
    return english_list, mandarin_list

def get_random_num_list(data_length, question_amount):
    return random.sample(range(data_length), question_amount)

def paper_and_answer_generator(english_list, mandarin_list, random_num_list):
    question_template = Template("$number. $question ____________")
    answer_template = Template("$number $question $answer ")

    paper = ""
    answer = ""
    count = 1
    for num in random_num_list:
        ques = question_template.substitute(number=count, question=mandarin_list[num])
        if len(ques) < 23:
            for _ in range(23 - len(ques)):
                ques += "_"
        ans = answer_template.substitute(number=count, question=mandarin_list[num], answer=english_list[num])
        if len(ans) < 20:
            for _ in range(20 - len(ans)):
                ans += "  "
        paper += ques
        answer += ans
        count += 1
        if count % 2 == 1:
            paper += "\n\n"
            answer += "\n\n"
        else:
            paper += "   "
            answer += "   "
    return paper, answer

def txt_to_docx(title, txt_file, ans_or_quiz):
    if ans_or_quiz == "ans":
        path = "./答案/"
    if ans_or_quiz == "quiz":
        path = "./考卷/"
    today = date.today().strftime ("%m/%d")
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(15)
    document.add_heading("Date: {0} Class:_____ Name:_______".format(today), 0)
    document.add_paragraph(txt_file, style=style)
    file_name = "{0}-{1}.{2}".format(title, date.today(), "docx")
    document.save(path + file_name)

if __name__ == "__main__":
    path, amount = get_user_input()
    eng_list, mandarin_list = load_data(path)
    random_num_list = get_random_num_list(len(eng_list), int(amount))
    question_paper, answer = paper_and_answer_generator(eng_list, mandarin_list, random_num_list)
    txt_to_docx("題目卷", question_paper, "quiz")
    txt_to_docx("答案卷", answer, "ans")